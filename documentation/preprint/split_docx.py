#!/usr/bin/env python3
"""
Split a combined Word document into main paper and supplement.
Properly handles all document elements (paragraphs, tables, etc.) in order.
"""

import sys
import os
from copy import deepcopy

try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx library not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)

def find_supplement_start(doc):
    """Find the paragraph index where supplement starts."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith("# KCOR: Supplementary Material") or text.startswith("## Supplementary material"):
            return i
    return len(doc.paragraphs)

def split_docx(combined_path, main_path, supplement_path):
    """Split combined Word document into main and supplement."""
    doc = Document(combined_path)
    
    # Find where supplement starts
    supplement_start_idx = find_supplement_start(doc)
    
    # Create new documents
    main_doc = Document()
    main_doc.core_properties.title = doc.core_properties.title or "KCOR"
    if doc.core_properties.author:
        main_doc.core_properties.author = doc.core_properties.author
    
    supp_doc = Document()
    supp_doc.core_properties.title = "KCOR: Supplementary Material"
    if doc.core_properties.author:
        supp_doc.core_properties.author = doc.core_properties.author
    
    # Get the body XML to iterate through elements in order
    body = doc._body._body
    main_body = main_doc._body._body
    supp_body = supp_doc._body._body
    
    para_count = 0
    table_count = 0
    
    # Iterate through all elements in document body in order
    for element in body:
        if element.tag == qn('w:p'):  # Paragraph
            if para_count < supplement_start_idx:
                # Add to main document
                main_body.append(deepcopy(element))
            else:
                # Add to supplement document
                supp_body.append(deepcopy(element))
            para_count += 1
        elif element.tag == qn('w:tbl'):  # Table
            # Determine which section based on current paragraph count
            if para_count <= supplement_start_idx:
                # Add to main document
                main_body.append(deepcopy(element))
            else:
                # Add to supplement document
                supp_body.append(deepcopy(element))
            table_count += 1
        else:
            # Other elements (sectPr, etc.) - add to both or handle appropriately
            # For now, add to main document
            if para_count < supplement_start_idx:
                main_body.append(deepcopy(element))
            else:
                supp_body.append(deepcopy(element))
    
    # Save documents
    main_doc.save(main_path)
    supp_doc.save(supplement_path)
    print(f"Split document: {combined_path}")
    print(f"  Main paper: {main_path} (split at paragraph {supplement_start_idx})")
    print(f"  Supplement: {supplement_path}")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: split_docx.py <combined.docx> <main.docx> <supplement.docx>", file=sys.stderr)
        sys.exit(1)
    
    combined_path = sys.argv[1]
    main_path = sys.argv[2]
    supplement_path = sys.argv[3]
    
    if not os.path.exists(combined_path):
        print(f"ERROR: Combined file not found: {combined_path}", file=sys.stderr)
        sys.exit(1)
    
    split_docx(combined_path, main_path, supplement_path)
